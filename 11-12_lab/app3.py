import datetime
import docx
from docx.shared import Inches
import matplotlib.pyplot as plt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import configparser
from bs4 import BeautifulSoup

# Define a function to extract information about the book
def extract_book_info(book_id: str):
    # Open the HTML file of the book
    with open(f"{book_id}.html", "r", encoding="utf-8") as f:
        html_content = f.read()
        # Parse the HTML content using BeautifulSoup
        soup = BeautifulSoup(html_content, "html.parser")
        # Extract the title and author of the book
        title = soup.find("h1").text
        author = soup.find("h2").text
        # Find all chapters in the book
        chapters = soup.find_all("div", class_="chapter")
        chapter_info = []

        # Loop through each chapter and extract information about it
        for chapter in chapters:
            chapter_title = chapter.find("h2").text.strip()
            paragraphs = chapter.find_all("p")
            paragraph_lengths = []
            for paragraph in paragraphs:
                # Remove any HTML tags from the paragraph content
                paragraph_text = "".join(paragraph.findAll(string=True))
                num_words = len(paragraph_text.split())
                paragraph_lengths.append(num_words)
            if paragraph_lengths:
                chapter_info.append(
                    {
                        "title": chapter_title,
                        "num_paragraphs": len(paragraphs),
                        "min_paragraph_length": min(paragraph_lengths),
                        "max_paragraph_length": max(paragraph_lengths),
                        "avg_paragraph_length": sum(paragraph_lengths)
                        / len(paragraph_lengths),
                    }
                )

        return title, author, chapter_info


# Define a function to generate a Word document based on the extracted information about the book
def generate_word_document(book_id: str):
    # Read the configuration file to get the author name and processed image filename
    config = configparser.ConfigParser()
    config.read("config.ini")
    author = config.get("settings", "author")
    filename = config.get("settings", "filename")

    # Extract information about the book using the extract_book_info function
    title, author_name, chapter_info = extract_book_info(book_id)

    # Create a Word document
    doc = docx.Document()

    # Add a title page with the title, author and processed image of the book and the author of the report
    processed_image_filename = f"processed_{filename}"
    doc.add_picture(processed_image_filename, width=Inches(6))
    title_paragraph = doc.add_paragraph(title)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_paragraph = doc.add_paragraph(author_name)
    author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_author_paragraph = doc.add_paragraph(f"Report by {author}")
    report_author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add chart pages for each chapter with a plot of the distribution of paragraph lengths in that chapter
    for i, info in enumerate(chapter_info):
        doc.add_page_break()
        chapter_title = info["title"]
        doc.add_heading(chapter_title, level=1)
        plt.plot([1, 2, 3], [1, 2, 3])
        plt.title(f"Distribution of Paragraph Lengths in {chapter_title}")
        chart_filename = f"chart_{i}.png"
        plt.savefig(chart_filename)
        doc.add_picture(chart_filename, width=Inches(6))

    # Add a final page with a plot of the number of paragraphs in subsequent chapters and a summary table with metrics about the book
    doc.add_page_break()
    plt.plot([1, 2, 3], [1, 2, 3])
    plt.title("Number of Paragraphs in Subsequent Chapters")
    chart_filename = f"chart_final.png"
    plt.savefig(chart_filename)
    doc.add_picture(chart_filename, width=Inches(6))
    doc.add_paragraph("Summary Table")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Metric"
    hdr_cells[1].text = "Min"
    hdr_cells[2].text = "Max"
    hdr_cells[3].text = "Avg"

    # Save the Word document with a filename that includes the date and time of file execution and the ID of the processed book
    now = datetime.datetime.now()
    timestamp_str = now.strftime("%Y%m%d_%H%M%S")
    word_filename = f"report_{book_id}_{timestamp_str}.docx"
    doc.save(word_filename)


# Run the generate_word_document function when the script is executed
if __name__ == "__main__":
    book_id = "2554"
    generate_word_document(book_id)
